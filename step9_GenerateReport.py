import pandas as pd
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

# ---------------- SETTINGS ----------------
# Output directory and report file path
make_dirs = True
if make_dirs:
    os.makedirs(r"D:\M.Tech DS\AIML application\Assignment-1\Step10_Report", exist_ok=True)
    
READ_DIR = r"D:\M.Tech DS\AIML application\Assignment-1\Step7_model_prediction"
OUTPUT_DIR = r"D:\M.Tech DS\AIML application\Assignment-1\Step10_Report"
REPORT_DOCX = os.path.join(OUTPUT_DIR, "Groundwater_Assignment_Report1.docx")

# Input files
coef_file = os.path.join(READ_DIR, "model_coefficients.csv")
summary_file = os.path.join(READ_DIR, "trained-model_summary.csv")
metrics_file = os.path.join(READ_DIR, "test-model-summary.csv")
interpret_file = os.path.join(READ_DIR, "Step10_Interpretation.csv")
pred_plot = os.path.join(READ_DIR, "actual_vs_predicted.png")
hist_plot = os.path.join(READ_DIR, "residuals_histogram.png")
qq_plot = os.path.join(READ_DIR, "residuals_qqplot.png")
resid_plot = os.path.join(READ_DIR, "residuals_vs_fitted.png")
residual_plot = os.path.join(READ_DIR, "residuals_vs_predicted.png")

# -------------------------------------------

# Load CSVs
coef_df = pd.read_csv(coef_file)
summary_df = pd.read_csv(summary_file)
metrics_df = pd.read_csv(metrics_file)
interpret_df = pd.read_csv(interpret_file)

# Remove fully NaN rows in interpretation
interpret_df = interpret_df.dropna(how='all', subset=['Section','Feature','Impact','Coef','P-value','Metric','Value','Description'])

# Create Word document
doc = Document()

# Set default font: Times New Roman, 12 pt
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
# Fix for Word XML
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

doc.add_heading("Groundwater Level Prediction Using Multiple Linear Regression", 0)

# ------------------ 1. Scenario ------------------
doc.add_heading("1. Scenario", level=1)
scenario_points = [
    "Delhi NCR",
    "Objective: Identify key drivers of groundwater depletion and predict groundwater levels at unsampled locations."
]
for pt in scenario_points:
    doc.add_paragraph(pt, style='List Bullet')

# ------------------ 2. Methodology ------------------
doc.add_heading("2. Methodology", level=1)
method_points = [
    "Data: Preprocessed groundwater dataset with 175 columns including the target 'data Value'.",
    "Dependent variable: dataValue",
    f"Independent variables: {len(coef_df)-1} selected features (BIC criterion)",
    "Spatial unit: District",
    "Temporal unit: Daily",
    "Model equation: Multiple Linear Regression (OLS)",
    "Data Acquisition: Data acquired from India WRIS, Bhuvan ISRO, Copernicus Climate Data Store, NICES Portal, SHRUG Atlas.",
    "Data Merging: Merged datasets on district and date.",
    "Data Preprocessing: Missing values handled, outliers removed, data merged appropriately.",
    "Feature Engineering: New features created based on domain knowledge.",
    "Model Specification: Defined model structure and selected features.",
    "Model Training: Trained the model using the training dataset.",
    "Model Evaluation: Evaluated model performance using test dataset and metrics like R², RMSE.",
    "Model Diagnostics: Residuals analyzed for patterns.",
    "EDA: Scatter plots, pair plots, and spatial-temporal plots analyzed to observe trends."
]
for pt in method_points:
    doc.add_paragraph(pt, style='List Bullet')

# ------------------ 3. EDA ------------------
doc.add_heading("3. Exploratory Data Analysis (EDA)", level=1)
eda_points = [
    "Scatter plots and pair plots were used to explore relationships between features and groundwater levels.",
    "Spatial and temporal trends were observed.",
    "Data was cleaned and outliers removed."
]
for pt in eda_points:
    doc.add_paragraph(pt, style='List Bullet')
if os.path.exists(pred_plot):
    doc.add_picture(pred_plot, width=Inches(5))

# ------------------ 4. Model Assumptions ------------------
doc.add_heading("4. Model Assumptions", level=1)
assump_points = [
    "Linearity: Relationships between predictors and target are linear.",
    "No Perfect Multicollinearity: Checked correlations among predictors.",
    "Exogeneity: Residuals uncorrelated with predictors.",
    "Homoscedasticity: Residuals have constant variance."
]
for pt in assump_points:
    doc.add_paragraph(pt, style='List Bullet')

# ------------------ 5. Model Selection ------------------
doc.add_heading("5. Model Selection", level=1)
doc.add_paragraph("Compared models using AIC and BIC. BIC-selected model (58 predictors) was chosen for analysis.", style='List Bullet')

# ------------------ 6. Model Estimation & Diagnostics ------------------
doc.add_heading("6. Model Estimation & Diagnostics", level=1)
doc.add_paragraph("Top 10 Coefficients:", style='List Bullet')
top_coef_df = coef_df.sort_values('p_value').head(10)
table = doc.add_table(rows=1, cols=len(top_coef_df.columns))
hdr_cells = table.rows[0].cells
for i, col in enumerate(top_coef_df.columns):
    hdr_cells[i].text = col
for _, row in top_coef_df.iterrows():
    cells = table.add_row().cells
    for i, col in enumerate(top_coef_df.columns):
        val = row[col]
        if isinstance(val, (float, int)):
            cells[i].text = str(round(val, 6))
        elif pd.notnull(val):
            cells[i].text = str(val)
        else:
            cells[i].text = ""

doc.add_paragraph("Model Fit Metrics:", style='List Bullet')
key_metrics = summary_df[summary_df['Metric'].isin(['R_squared(training)','Adj_R_squared','F_stat','F_pvalue','Num_Predictors' ])]
for _, row in key_metrics.iterrows():
    doc.add_paragraph(f"{row['Metric']}: {row['Value']:.4f}", style='List Bullet')

# ------------------ 7. Predictions & Evaluation ------------------
doc.add_heading("7. Predictions & Evaluation", level=1)
for _, row in metrics_df.iterrows():
    doc.add_paragraph(f"{row['Metric']}: {row['Value']:.4f}", style='List Bullet')
doc.add_paragraph("Plots:", style='List Bullet')
if os.path.exists(pred_plot):
    doc.add_picture(pred_plot, width=Inches(5))
if os.path.exists(hist_plot):
    doc.add_picture(hist_plot, width=Inches(5))
if os.path.exists(qq_plot):
    doc.add_picture(qq_plot, width=Inches(5))
if os.path.exists(resid_plot):
    doc.add_picture(resid_plot, width=Inches(5))
if os.path.exists(residual_plot):
    doc.add_picture(residual_plot, width=Inches(5))
    doc.add_paragraph("Residual Plot", style='List Bullet')
    doc.add_paragraph("The residual plot shows the difference between observed and predicted values. It helps in diagnosing the model fit.", style='List Bullet')
    doc.add_paragraph("Interpretation: No clear pattern suggests a good fit.", style='List Bullet')
    doc.add_paragraph("Action: Consider refining the model if patterns are detected.", style='List Bullet')

# ------------------ 8. Significant Features & Interpretation ------------------
doc.add_heading("8. Significant Features & Interpretation", level=1)
for _, row in interpret_df.iterrows():
    parts = []
    for col in ['Section','Feature','Impact','Coef','P-value','Metric','Value','Description']:
        if pd.notnull(row.get(col)):
            parts.append(f"{col}: {row[col]}")
    doc.add_paragraph("; ".join(parts), style='List Bullet')

# ------------------ 9. Conclusion & Policy Implications ------------------
doc.add_heading("9. Conclusion & Policy Implications", level=1)
conclusion_points = [
    "The model identifies key factors affecting groundwater levels.",
    "Predictions can guide water resource planning.",
    "Recommendations: Monitor significant drivers and use the model for short-term planning and risk assessment."
]
for pt in conclusion_points:
    doc.add_paragraph(pt, style='List Bullet')

# ------------------ 10. References ------------------
doc.add_heading("10. References", level=1)
refs_points = [
    "India WRIS: https://indiawris.gov.in/wris/",
    "Bhuvan ISRO: https://bhuvan-app1.nrsc.gov.in/2dresources/bhuvanstore2.php",
    "Copernicus Climate Data Store: https://cds.climate.copernicus.eu/#!/home",
    "NICES Portal: https://nices.nrsc.gov.in/",
    "SHRUG Atlas: https://www.devdatalab.org/atlas"
]
for pt in refs_points:
    doc.add_paragraph(pt, style='List Bullet')

# ------------------ Save Document ------------------
doc.save(REPORT_DOCX)
print(f"Report generated and saved at: {REPORT_DOCX}")